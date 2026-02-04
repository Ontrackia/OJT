"""
Knowledge Ingestion Service - MOE and Procedures RAG
Indexes company documents with priority over general regulations
"""
from sqlalchemy.orm import Session
from typing import Optional, List, BinaryIO
from datetime import datetime
import hashlib
import os
import uuid
from PyPDF2 import PdfReader
import docx
from cryptography.fernet import Fernet
import chromadb
from sentence_transformers import SentenceTransformer

from app.models.evidence_vault import KnowledgeDocument


class KnowledgeIngestionService:
    """
    MOE and Procedures Ingestion Service
    
    Features:
    - PDF/DOCX parsing and chunking
    - ChromaDB indexing with priority
    - RAG retrieval with MOE prioritization
    - Version control
    """
    
    CHUNK_SIZE = 500  # Characters per chunk
    CHUNK_OVERLAP = 50
    ENCRYPTION_KEY = os.getenv("EVIDENCE_ENCRYPTION_KEY")
    
    def __init__(self, db: Session, organization_id: int):
        self.db = db
        self.organization_id = organization_id
        
        # Initialize encryption
        if not self.ENCRYPTION_KEY:
            raise ValueError("EVIDENCE_ENCRYPTION_KEY not set")
        self.cipher = Fernet(self.ENCRYPTION_KEY.encode())
        
        # Initialize ChromaDB
        self.chroma_client = chromadb.PersistentClient(path="./chroma_db")
        self.embedding_model = SentenceTransformer('all-MiniLM-L6-v2')
        
        # Get or create collection for this organization
        self.collection_name = f"org_{organization_id}_knowledge"
        self.collection = self.chroma_client.get_or_create_collection(
            name=self.collection_name,
            metadata={"organization_id": organization_id}
        )
    
    def ingest_document(
        self,
        file: BinaryIO,
        filename: str,
        document_type: str,  # MOE, QUALITY_MANUAL, PROCEDURE, SOP
        title: str,
        description: Optional[str] = None,
        uploaded_by: int = 1,
        uploaded_by_name: str = "Admin",
        version: str = "1.0"
    ) -> KnowledgeDocument:
        """
        Ingest MOE or procedure document.
        
        Steps:
        1. Extract text from PDF/DOCX
        2. Chunk text intelligently
        3. Generate embeddings
        4. Store in ChromaDB with priority metadata
        5. Encrypt and save original file
        6. Create database record
        """
        # Read file
        file_content = file.read()
        file_hash = hashlib.sha256(file_content).hexdigest()
        
        # Extract text
        extension = os.path.splitext(filename)[1].lower()
        text_content = self._extract_text(file_content, extension)
        
        # Chunk text
        chunks = self._chunk_text(text_content)
        
        print(f"📄 Extracted {len(text_content)} characters, created {len(chunks)} chunks")
        
        # Generate document ID
        document_id = f"DOC-{datetime.utcnow().strftime('%Y%m%d')}-{str(uuid.uuid4())[:8].upper()}"
        
        # Determine priority (MOE = highest)
        priority = self._get_priority(document_type)
        
        # Index in ChromaDB
        chunk_ids = []
        for i, chunk in enumerate(chunks):
            chunk_id = f"{document_id}-chunk-{i}"
            chunk_ids.append(chunk_id)
            
            # Generate embedding
            embedding = self.embedding_model.encode(chunk).tolist()
            
            # Add to ChromaDB with metadata
            self.collection.add(
                ids=[chunk_id],
                embeddings=[embedding],
                documents=[chunk],
                metadatas=[{
                    "document_id": document_id,
                    "document_type": document_type,
                    "title": title,
                    "chunk_index": i,
                    "priority": priority,
                    "organization_id": self.organization_id,
                    "version": version
                }]
            )
        
        # Encrypt original file
        encrypted_content = self.cipher.encrypt(file_content)
        encrypted_path = self._save_encrypted_document(
            encrypted_content,
            document_id,
            extension
        )
        
        # Create database record
        knowledge_doc = KnowledgeDocument(
            tenant_id=1,  # TODO: Get from context
            organization_id=self.organization_id,
            document_id=document_id,
            document_type=document_type,
            title=title,
            description=description,
            original_filename=filename,
            encrypted_file_path=encrypted_path,
            file_hash_sha256=file_hash,
            is_indexed=True,
            indexed_at=datetime.utcnow(),
            chunk_count=len(chunks),
            vector_collection_id=self.collection_name,
            priority=priority,
            version=version,
            uploaded_by=uploaded_by,
            uploaded_by_name=uploaded_by_name,
            retention_until=datetime.utcnow().replace(year=datetime.utcnow().year + 5)
        )
        
        self.db.add(knowledge_doc)
        self.db.commit()
        self.db.refresh(knowledge_doc)
        
        print(f"✅ Knowledge document ingested: {document_id} | {len(chunks)} chunks | Priority: {priority}")
        
        return knowledge_doc
    
    def query_knowledge(
        self,
        question: str,
        n_results: int = 5,
        territory: Optional[str] = None
    ) -> dict:
        """
        Query knowledge base with MOE prioritization.
        
        Returns results prioritized by:
        1. MOE (priority 100)
        2. Quality Manuals (priority 90)
        3. Procedures (priority 80)
        4. General regulations (priority 50)
        """
        # Generate query embedding
        query_embedding = self.embedding_model.encode(question).tolist()
        
        # Query ChromaDB
        results = self.collection.query(
            query_embeddings=[query_embedding],
            n_results=n_results * 2,  # Get more to filter by priority
            where={"organization_id": self.organization_id}
        )
        
        # Sort by priority
        sorted_results = []
        for i, doc_id in enumerate(results['ids'][0]):
            metadata = results['metadatas'][0][i]
            sorted_results.append({
                "chunk_id": doc_id,
                "document_id": metadata['document_id'],
                "document_type": metadata['document_type'],
                "title": metadata['title'],
                "content": results['documents'][0][i],
                "priority": metadata['priority'],
                "distance": results['distances'][0][i] if 'distances' in results else 0
            })
        
        # Sort by priority (descending) then distance (ascending)
        sorted_results.sort(key=lambda x: (-x['priority'], x['distance']))
        
        # Take top n_results
        top_results = sorted_results[:n_results]
        
        # Format response with citations
        context = ""
        citations = []
        
        for i, result in enumerate(top_results):
            source = f"{result['document_type']}: {result['title']}"
            context += f"\n[{i+1}] Según su {result['document_type']}, {result['title']}:\n{result['content']}\n"
            citations.append({
                "index": i+1,
                "source": source,
                "document_id": result['document_id'],
                "priority": result['priority']
            })
        
        return {
            "context": context,
            "citations": citations,
            "results": top_results
        }
    
    def _extract_text(self, content: bytes, extension: str) -> str:
        """Extract text from PDF or DOCX"""
        if extension == '.pdf':
            return self._extract_pdf_text(content)
        elif extension in ['.docx', '.doc']:
            return self._extract_docx_text(content)
        elif extension == '.txt':
            return content.decode('utf-8')
        else:
            raise ValueError(f"Unsupported file type: {extension}")
    
    def _extract_pdf_text(self, content: bytes) -> str:
        """Extract text from PDF"""
        import io
        pdf_file = io.BytesIO(content)
        reader = PdfReader(pdf_file)
        
        text = ""
        for page in reader.pages:
            text += page.extract_text() + "\n"
        
        return text
    
    def _extract_docx_text(self, content: bytes) -> str:
        """Extract text from DOCX"""
        import io
        doc_file = io.BytesIO(content)
        doc = docx.Document(doc_file)
        
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        return text
    
    def _chunk_text(self, text: str) -> List[str]:
        """
        Chunk text intelligently.
        
        Strategy:
        - Split by paragraphs first
        - Combine small paragraphs
        - Split large paragraphs
        - Maintain overlap for context
        """
        paragraphs = text.split('\n\n')
        chunks = []
        current_chunk = ""
        
        for para in paragraphs:
            para = para.strip()
            if not para:
                continue
            
            if len(current_chunk) + len(para) < self.CHUNK_SIZE:
                current_chunk += para + "\n\n"
            else:
                if current_chunk:
                    chunks.append(current_chunk.strip())
                current_chunk = para + "\n\n"
        
        if current_chunk:
            chunks.append(current_chunk.strip())
        
        return chunks
    
    def _get_priority(self, document_type: str) -> int:
        """Get priority for document type"""
        priorities = {
            "MOE": 100,
            "QUALITY_MANUAL": 90,
            "PROCEDURE": 80,
            "SOP": 80,
            "REGULATION": 50
        }
        return priorities.get(document_type, 50)
    
    def _save_encrypted_document(
        self,
        encrypted_content: bytes,
        document_id: str,
        extension: str
    ) -> str:
        """Save encrypted document to disk"""
        storage_path = f"./knowledge_vault/{self.organization_id}"
        os.makedirs(storage_path, exist_ok=True)
        
        filename = f"{document_id}{extension}.enc"
        file_path = os.path.join(storage_path, filename)
        
        with open(file_path, 'wb') as f:
            f.write(encrypted_content)
        
        return file_path
    
    def list_documents(
        self,
        document_type: Optional[str] = None,
        is_active: bool = True
    ) -> List[KnowledgeDocument]:
        """List knowledge documents"""
        query = self.db.query(KnowledgeDocument).filter(
            KnowledgeDocument.organization_id == self.organization_id,
            KnowledgeDocument.is_active == is_active
        )
        
        if document_type:
            query = query.filter(KnowledgeDocument.document_type == document_type)
        
        return query.order_by(KnowledgeDocument.uploaded_at.desc()).all()
